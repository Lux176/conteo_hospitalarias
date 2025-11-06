import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import plotly.express as px
import tempfile
import os
from datetime import datetime
from io import BytesIO
import unicodedata
import base64
import re

# Configuración de la página
st.set_page_config(
    page_title="Analizador de Incidentes",
    page_icon="📊",
    layout="wide"
)

# --- FUNCIONES MEJORADAS Y CORREGIDAS ---

def limpiar_texto(texto):
    """Limpia texto: minúsculas, sin acentos - MEJORADA"""
    if not isinstance(texto, str):
        return texto
    try:
        # Manejar mejor caracteres especiales
        texto = str(texto).strip()
        texto_limpio = unicodedata.normalize('NFD', texto)
        texto_limpio = texto_limpio.encode('ascii', 'ignore').decode('utf-8')
        texto_limpio = texto_limpio.lower().strip()
        return texto_limpio
    except (UnicodeError, AttributeError):
        # Fallback: convertir a string y limpiar básicamente
        return str(texto).strip().lower()

def parsear_fecha(fecha):
    """Parsea fechas en diferentes formatos - MEJORADA"""
    if pd.isna(fecha) or fecha is None or fecha == '':
        return None
    if isinstance(fecha, (datetime, pd.Timestamp)):
        return fecha
    
    fecha_str = str(fecha).strip()
    # Limpiar posibles espacios extra
    fecha_str = re.sub(r'\s+', ' ', fecha_str)
    
    formatos = [
        '%d/%m/%Y', '%d/%m/%y', '%Y-%m-%d', '%d-%m-%Y',
        '%m/%d/%Y', '%m/%d/%y', '%Y/%m/%d', '%d.%m.%Y',
        '%d %m %Y', '%d %b %Y', '%d %B %Y'
    ]
    
    for fmt in formatos:
        try:
            parsed = datetime.strptime(fecha_str, fmt)
            # Validar que sea una fecha razonable (no futura más de 1 año, no anterior a 2000)
            if parsed.year >= 2000 and parsed <= datetime.now().replace(year=datetime.now().year + 1):
                return parsed
        except ValueError:
            continue
    
    return None

def generar_grafica_bar(conteo, titulo, filename):
    """Genera gráficas usando matplotlib - MEJORADA"""
    if conteo.empty:
        # Crear gráfica vacía si no hay datos
        plt.figure(figsize=(8, 4))
        plt.text(0.5, 0.5, 'No hay datos para mostrar', 
                ha='center', va='center', transform=plt.gca().transAxes, fontsize=12)
        plt.title(titulo)
    else:
        df_plot = conteo.reset_index()
        df_plot.columns = ['Tipo de Incidente', 'Cantidad']
        
        # Limitar a los top 20 si hay muchos tipos
        if len(df_plot) > 20:
            df_plot = df_plot.nlargest(20, 'Cantidad')
            titulo += " (Top 20)"
        
        plt.figure(figsize=(12, 6))
        colors = plt.cm.viridis(np.linspace(0, 1, len(df_plot)))
        bars = plt.bar(df_plot['Tipo de Incidente'], df_plot['Cantidad'], color=colors)
        
        plt.title(titulo, fontsize=14, fontweight='bold')
        plt.xlabel('Tipo de Incidente', fontweight='bold')
        plt.ylabel('Cantidad', fontweight='bold')
        plt.xticks(rotation=45, ha='right')
        
        # Añadir valores en las barras
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{int(height)}', ha='center', va='bottom', fontweight='bold')
        
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
    
    # Guardar imagen
    path = os.path.join(tempfile.gettempdir(), filename)
    plt.savefig(path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return path

def generar_grafica_plotly(conteo, titulo):
    """Genera gráfica plotly para mostrar en Streamlit - MEJORADA"""
    if conteo.empty:
        # Crear gráfica vacía
        fig = px.bar(title=titulo + " - No hay datos")
        fig.update_layout(
            annotations=[dict(text="No hay datos para mostrar", x=0.5, y=0.5, showarrow=False)]
        )
    else:
        df_plot = conteo.reset_index()
        df_plot.columns = ['Tipo de Incidente', 'Cantidad']
        
        # Limitar a top 20 si hay muchos tipos
        if len(df_plot) > 20:
            df_plot = df_plot.nlargest(20, 'Cantidad')
            titulo += " (Top 20)"
        
        fig = px.bar(df_plot, x='Tipo de Incidente', y='Cantidad', title=titulo,
                     color='Cantidad', color_continuous_scale='Viridis')
        fig.update_layout(
            xaxis_tickangle=-45,
            plot_bgcolor='white',
            paper_bgcolor='white',
            font=dict(size=12)
        )
        fig.update_xaxes(title_text="Tipo de Incidente")
        fig.update_yaxes(title_text="Cantidad")
    
    return fig

def es_traslado_afirmativo(valor):
    """Determina si un valor indica traslado afirmativo - NUEVA FUNCIÓN"""
    if pd.isna(valor) or valor is None:
        return False
    
    valor_str = str(valor).strip().lower()
    
    # Lista más completa de valores afirmativos
    afirmativos = [
        'sí', 'si', 's', 'yes', 'true', '1', 'verdadero', 
        'afirmativo', 'x', '✓', 'check', 'checked', 'ok',
        'aceptar', 'confirmado', 'realizado'
    ]
    
    # También considerar números mayores a 0
    try:
        if float(valor_str) > 0:
            return True
    except (ValueError, TypeError):
        pass
    
    return valor_str in afirmativos

def generar_reporte_word(conteos, traslados_info, imagenes):
    """Genera reporte en formato Word - MEJORADA"""
    doc = Document()
    
    # Título principal
    title = doc.add_heading('Análisis de Incidentes', 0)
    title.alignment = 1  # Centrado
    
    # Fecha de generación
    doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()
    
    # Resumen de conteos
    doc.add_heading('Resumen de Conteos', level=1)
    for nombre, conteo in conteos.items():
        doc.add_heading(nombre, level=2)
        if not conteo.empty:
            doc.add_paragraph(f"Total de incidentes: {conteo.sum()}")
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = "Tipo de Incidente"
            hdr_cells[1].text = "Cantidad"
            
            for tipo, cantidad in conteo.items():
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(tipo)
                row_cells[1].text = str(cantidad)
        else:
            doc.add_paragraph("No hay datos disponibles")

    # Resumen de traslados
    if traslados_info:
        doc.add_heading('Resumen de Traslados', level=1)
        for k, v in traslados_info.items():
            doc.add_paragraph(f"{k}: {v}", style='List Bullet')

    # Gráficas
    doc.add_page_break()
    doc.add_heading('Gráficas', level=1)
    for titulo, path in imagenes.items():
        if os.path.exists(path):
            doc.add_heading(titulo, level=2)
            try:
                doc.add_picture(path, width=Inches(5.5))
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"Error al cargar gráfica: {str(e)}")

    output_path = os.path.join(tempfile.gettempdir(), f'reporte_incidentes_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(output_path)
    return output_path

def generar_reporte_txt(conteos, traslados_info):
    """Genera reporte en formato de texto - MEJORADA"""
    texto = ["Análisis de Incidentes"]
    texto.append("=" * 50)
    texto.append(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    texto.append("\nRESUMEN DE CONTEOS")
    texto.append("=" * 50)
    
    for nombre, conteo in conteos.items():
        texto.append(f"\n{nombre.upper()}")
        texto.append("-" * len(nombre))
        if not conteo.empty:
            texto.append(f"Total de incidentes: {conteo.sum()}")
            for tipo, cantidad in conteo.items():
                texto.append(f"  {tipo}: {cantidad}")
        else:
            texto.append("No hay datos disponibles")
    
    if traslados_info:
        texto.append("\n\nRESUMEN DE TRASLADOS")
        texto.append("-" * 20)
        for k, v in traslados_info.items():
            texto.append(f"{k}: {v}")
    
    contenido = "\n".join(texto)
    path_txt = os.path.join(tempfile.gettempdir(), f"reporte_incidentes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
    with open(path_txt, "w", encoding="utf-8") as f:
        f.write(contenido)
    return path_txt

def get_download_link(file_path, file_label):
    """Genera un enlace de descarga para el archivo - MEJORADA"""
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        b64 = base64.b64encode(data).decode()
        file_name = os.path.basename(file_path)
        href = f'<a href="data:file/octet-stream;base64,{b64}" download="{file_name}" style="background-color: #4CAF50; color: white; padding: 10px 20px; text-align: center; text-decoration: none; display: inline-block; border-radius: 5px; margin: 5px;">📥 {file_label}</a>'
        return href
    except Exception as e:
        return f'<p style="color: red;">Error al generar enlace: {str(e)}</p>'

# --- INTERFAZ PRINCIPAL MEJORADA ---

def main():
    st.title("📊 Analizador de Incidentes")
    st.markdown("---")
    
    # Carga de archivo
    st.header("1. Carga de Datos")
    uploaded_file = st.file_uploader("Sube tu archivo de datos (CSV o Excel)", type=['csv', 'xlsx'])
    
    if uploaded_file is not None:
        try:
            # Leer archivo
            if uploaded_file.name.endswith('.xlsx'):
                df = pd.read_excel(uploaded_file)
            else:
                df = pd.read_csv(uploaded_file, encoding='utf-8', on_bad_lines='skip')
            
            st.success(f"✅ Archivo cargado correctamente. Dimensiones: {df.shape[0]} filas × {df.shape[1]} columnas")
            
            # Mostrar vista previa
            with st.expander("📋 Vista previa de los datos"):
                st.dataframe(df.head(10))
                st.write(f"**Tipos de datos:**")
                st.write(df.dtypes)
            
            # Selección de columnas
            st.header("2. Configuración del Análisis")
            
            col1, col2 = st.columns(2)
            with col1:
                col_incidentes = st.selectbox(
                    "Selecciona la columna de INCIDENTES:",
                    options=df.columns,
                    index=next((i for i, col in enumerate(df.columns) if 'incident' in col.lower()), 0) if any('incident' in col.lower() for col in df.columns) else None
                )
            
            with col2:
                col_traslado = st.selectbox(
                    "Selecciona la columna de TRASLADO A HOSPITAL:",
                    options=df.columns,
                    index=next((i for i, col in enumerate(df.columns) if any(word in col.lower() for word in ['traslado', 'hospital', 'transfer']), 0) if any(any(word in col.lower() for word in ['traslado', 'hospital', 'transfer']) for col in df.columns) else None
                )
            
            # Filtro por fechas - CORREGIDO
            st.subheader("🗓️ Filtro por Fechas (Opcional)")
            usar_fechas = st.checkbox("Activar filtro por fechas")
            
            fecha_inicio = None
            fecha_fin = None
            col_fechas = None

            if usar_fechas:  # ← CORREGIDO: Removido "and col_incidentes"
                col_fechas = st.selectbox(
                    "Selecciona la columna de FECHAS:",
                    options=df.columns,
                    index=next((i for i, col in enumerate(df.columns) if any(word in col.lower() for word in ['fecha', 'date']), 0) if any(any(word in col.lower() for word in ['fecha', 'date']) for col in df.columns) else None
                )
                
                if col_fechas:
                    col3, col4 = st.columns(2)
                    with col3:
                        fecha_inicio_str = st.text_input("Fecha de inicio (d/m/AAAA):", placeholder="01/01/2024")
                    with col4:
                        fecha_fin_str = st.text_input("Fecha de fin (d/m/AAAA):", placeholder="31/12/2024")
                    
                    if fecha_inicio_str and fecha_fin_str:
                        try:
                            fecha_inicio = datetime.strptime(fecha_inicio_str.strip(), '%d/%m/%Y')
                            fecha_fin = datetime.strptime(fecha_fin_str.strip(), '%d/%m/%Y')
                            
                            if fecha_inicio > fecha_fin:
                                st.error("❌ La fecha de inicio no puede ser mayor que la fecha de fin")
                            else:
                                st.info(f"📅 Rango seleccionado: {fecha_inicio.strftime('%d/%m/%Y')} - {fecha_fin.strftime('%d/%m/%Y')}")
                                
                        except ValueError:
                            st.error("❌ Formato de fecha incorrecto. Use el formato d/m/AAAA (ej: 01/01/2024)")
            
            # Análisis por Servicios Médicos - MEJORADO
            st.subheader("🏥 Análisis por Servicios Médicos")
            incluir_sm = st.checkbox("Incluir análisis por Servicios Médicos")
            
            col_sm = None
            valor_sm = 'SM'  # Valor por defecto
            
            if incluir_sm:
                col_sm = st.selectbox(
                    "Selecciona la columna de SERVICIOS MÉDICOS:",
                    options=df.columns,
                    index=next((i for i, col in enumerate(df.columns) if any(word in col.lower() for word in ['servicio', 'medic', 'sm']), 0) if any(any(word in col.lower() for word in ['servicio', 'medic', 'sm']) for col in df.columns) else None
                )
                
                if col_sm:
                    # Mostrar valores únicos para ayudar al usuario
                    with st.expander("🔍 Ver valores únicos en esta columna"):
                        valores_unicos = df[col_sm].dropna().unique()[:20]  # Mostrar primeros 20
                        st.write(f"Valores encontrados: {list(valores_unicos)}")
                    
                    valor_sm = st.text_input(
                        "Valor que identifica Servicios Médicos:",
                        value="SM",
                        help="Ingresa el valor exacto que identifica los Servicios Médicos en tu datos"
                    )
            
            # Botón para generar análisis
            if col_incidentes and col_traslado:
                st.markdown("---")
                if st.button("🚀 Generar Reporte Completo", type="primary", use_container_width=True):
                    
                    with st.spinner("Procesando datos..."):
                        # Crear copia para no modificar el original
                        df_clean = df.copy()
                        
                        # Aplicar filtro de fechas si está activado
                        if usar_fechas and fecha_inicio and fecha_fin and col_fechas:
                            df_clean['fecha_parseada'] = df_clean[col_fechas].apply(parsear_fecha)
                            df_filtrado = df_clean.dropna(subset=['fecha_parseada'])
                            df_filtrado = df_filtrado[
                                (df_filtrado['fecha_parseada'] >= fecha_inicio) & 
                                (df_filtrado['fecha_parseada'] <= fecha_fin)
                            ]
                            st.info(f"📊 Datos filtrados: {len(df_filtrado)} registros de {len(df_clean)} originales")
                            df_clean = df_filtrado
                        
                        # Verificar que hay datos después del filtrado
                        if df_clean.empty:
                            st.error("❌ No hay datos después del filtrado. Ajusta los criterios de filtro.")
                            return
                        
                        # Generar conteos y traslados
                        conteos = {}
                        traslados_info = {}
                        
                        # Conteo general
                        if col_incidentes in df_clean.columns:
                            conteos["Total de Incidentes"] = df_clean[col_incidentes].value_counts()
                        
                        # Traslados generales - USANDO FUNCIÓN MEJORADA
                        if col_traslado in df_clean.columns:
                            traslados_info["Total de traslados"] = df_clean[col_traslado].apply(es_traslado_afirmativo).sum()
                        
                        # Análisis por Servicios Médicos - MEJORADO
                        if incluir_sm and col_sm and col_sm in df_clean.columns:
                            # Limpiar columna SM
                            df_clean[col_sm] = df_clean[col_sm].astype(str).str.strip().str.upper()
                            
                            # Usar el valor proporcionado por el usuario
                            df_sm = df_clean[df_clean[col_sm] == valor_sm.upper()]
                            df_no_sm = df_clean[df_clean[col_sm] != valor_sm.upper()]
                            
                            if not df_sm.empty and col_incidentes in df_sm.columns:
                                conteos["Incidentes atendidos por Servicios Médicos"] = df_sm[col_incidentes].value_counts()
                            
                            if not df_no_sm.empty and col_incidentes in df_no_sm.columns:
                                conteos["Incidentes atendidos por Operativa Médica Protección Civil"] = df_no_sm[col_incidentes].value_counts()
                            
                            # Traslados por categoría
                            if col_traslado in df_sm.columns:
                                traslados_info["Traslados por Servicios Médicos"] = df_sm[col_traslado].apply(es_traslado_afirmativo).sum()
                            
                            if col_traslado in df_no_sm.columns:
                                traslados_info["Traslados por Operativa Médica Protección Civil"] = df_no_sm[col_traslado].apply(es_traslado_afirmativo).sum()
                        
                        # Mostrar resultados en la interfaz
                        st.header("3. 📈 Resultados del Análisis")
                        
                        # Métricas rápidas
                        col_met1, col_met2, col_met3 = st.columns(3)
                        with col_met1:
                            total_incidentes = len(df_clean)
                            st.metric("Total de Incidentes", total_incidentes)
                        with col_met2:
                            tipos_incidentes = len(conteos.get("Total de Incidentes", pd.Series()))
                            st.metric("Tipos de Incidentes", tipos_incidentes)
                        with col_met3:
                            total_traslados = traslados_info.get("Total de traslados", 0)
                            st.metric("Total de Traslados", total_traslados)
                        
                        # Mostrar información de traslados
                        if traslados_info:
                            st.subheader("🚑 Resumen de Traslados")
                            for k, v in traslados_info.items():
                                col_t1, col_t2 = st.columns([2, 1])
                                with col_t1:
                                    st.write(f"**{k}**")
                                with col_t2:
                                    st.write(f"**{v}**")
                        
                        # Mostrar tablas y gráficas de conteos
                        for nombre, conteo in conteos.items():
                            st.subheader(nombre)
                            
                            if not conteo.empty:
                                # Mostrar tabla
                                df_display = conteo.reset_index()
                                df_display.columns = ['Tipo de Incidente', 'Cantidad']
                                st.dataframe(df_display, use_container_width=True)
                                
                                # Mostrar gráfica interactiva
                                fig = generar_grafica_plotly(conteo, nombre)
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info("No hay datos disponibles para este análisis")
                        
                        # Generar gráficas para el reporte Word
                        st.header("4. 📄 Generando Reportes Descargables")
                        with st.spinner("Generando gráficas para el reporte..."):
                            imagenes = {}
                            for k, v in conteos.items():
                                if not v.empty:
                                    safe_filename = f"grafica_{k.replace(' ', '_').replace('/', '_').lower()}_{datetime.now().strftime('%H%M%S')}.png"
                                    imagenes[k] = generar_grafica_bar(v, k, safe_filename)
                        
                        # Generar y ofrecer descarga de reportes
                        st.success("✅ Reportes generados correctamente")
                        
                        col_dl1, col_dl2 = st.columns(2)
                        
                        with col_dl1:
                            with st.spinner("Generando reporte Word..."):
                                try:
                                    doc_path = generar_reporte_word(conteos, traslados_info, imagenes)
                                    st.markdown(get_download_link(doc_path, "Descargar Reporte Word (.docx)"), unsafe_allow_html=True)
                                except Exception as e:
                                    st.error(f"❌ Error al generar reporte Word: {str(e)}")
                        
                        with col_dl2:
                            with st.spinner("Generando reporte de texto..."):
                                try:
                                    txt_path = generar_reporte_txt(conteos, traslados_info)
                                    st.markdown(get_download_link(txt_path, "Descargar Reporte Texto (.txt)"), unsafe_allow_html=True)
                                except Exception as e:
                                    st.error(f"❌ Error al generar reporte texto: {str(e)}")
                        
                        # Limpiar archivos temporales
                        for path in imagenes.values():
                            try:
                                if os.path.exists(path):
                                    os.remove(path)
                            except:
                                pass
            
            else:
                st.warning("⚠️ Por favor, selecciona las columnas de INCIDENTES y TRASLADO para continuar.")
                
        except Exception as e:
            st.error(f"❌ Error al procesar el archivo: {str(e)}")
            st.info("💡 **Sugerencias para solucionar el problema:**")
            st.info("- Verifica que el archivo no esté corrupto")
            st.info("- Asegúrate de que el archivo tenga el formato correcto")
            st.info("- Si es CSV, verifica la codificación (intenta UTF-8)")
            st.info("- Si es Excel, verifica que no tenga protecciones o formatos especiales")
    
    else:
        st.info("👆 Por favor, sube un archivo CSV o Excel para comenzar el análisis.")
        st.markdown("""
        ### 📝 Instrucciones:
        1. **Sube tu archivo** de datos (CSV o Excel)
        2. **Selecciona las columnas** correspondientes
        3. **Configura los filtros** si es necesario  
        4. **Genera el reporte** y descarga los resultados
        
        ### 🔍 Análisis que se generan:
        - Conteo general de incidentes
        - Análisis de traslados a hospital
        - Segmentación por Servicios Médicos (opcional)
        - Gráficas y reportes descargables
        """)

if __name__ == "__main__":
    main()
